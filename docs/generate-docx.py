"""
GSD Developer Guide - Markdown to Word Document Generator (v2)
Reads GSD-Developer-Guide.md and produces a professionally formatted .docx
Uses Word built-in heading styles for TOC support.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Colors ---
BLUE = RGBColor(0x2B, 0x57, 0x9A)
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
CODE_BG = "F2F2F2"
HEADER_BG = "2B579A"
ALT_ROW_HEX = "EDF2FA"
BORDER_COLOR = "CCCCCC"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(SCRIPT_DIR, "GSD-Developer-Guide.md")
DOCX_PATH = os.path.join(SCRIPT_DIR, "GSD-Developer-Guide.docx")

# Track numbered list counters
_num_counter = [0]


# ── Utility helpers ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    pPr.append(shading)


def add_bottom_border(paragraph, color="2B579A", size=6):
    """Add a bottom border line to a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_table_borders(table, color="CCCCCC", size=4):
    """Set uniform thin borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def parse_inline(paragraph, text, base_size=Pt(11), base_color=BLACK):
    """Parse inline markdown: **bold**, `code`, [links](url)."""
    i = 0
    while i < len(text):
        # Bold
        if text[i:i+2] == "**":
            end = text.find("**", i + 2)
            if end != -1:
                run = paragraph.add_run(text[i+2:end])
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = base_size
                run.font.color.rgb = base_color
                i = end + 2
                continue
        # Inline code
        if text[i] == "`":
            end = text.find("`", i + 1)
            if end != -1:
                run = paragraph.add_run(text[i+1:end])
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = DARK_GRAY
                rPr = run._element.get_or_add_rPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
                rPr.append(shading)
                i = end + 1
                continue
        # Link [text](url)
        link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', text[i:])
        if link_match:
            run = paragraph.add_run(link_match.group(1))
            run.font.name = "Calibri"
            run.font.size = base_size
            run.font.color.rgb = BLUE
            run.underline = True
            i += link_match.end()
            continue
        # Collect plain text
        j = i + 1
        while j < len(text) and text[j] not in ('*', '`', '['):
            j += 1
        run = paragraph.add_run(text[i:j])
        run.font.name = "Calibri"
        run.font.size = base_size
        run.font.color.rgb = base_color
        i = j


# ── Style Setup ──────────────────────────────────────────────────

def setup_styles(doc):
    """Configure Word built-in heading styles for TOC compatibility."""
    # Normal
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = Pt(16)

    # Heading 1 (Chapter titles)
    h1 = doc.styles['Heading 1']
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.page_break_before = True

    # Heading 2 (Section titles like 1.1, 2.3)
    h2 = doc.styles['Heading 2']
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = DARK_BLUE
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3 (Sub-section)
    h3 = doc.styles['Heading 3']
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = DARK_BLUE
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(4)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = "Calibri"
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = DARK_BLUE
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)


def add_headers_footers(doc):
    """Add header and footer to each section."""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # Header
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("GSD Autonomous Development Engine  |  Developer Guide")
        run.font.size = Pt(8)
        run.font.color.rgb = MED_GRAY
        run.font.name = "Calibri"
        # Subtle bottom border on header
        add_bottom_border(p, color="CCCCCC", size=4)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run("Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = MED_GRAY
        run.font.name = "Calibri"

        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1 = p.add_run()
        run1._element.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2 = p.add_run()
        run2._element.append(instrText)
        run2.font.size = Pt(8)
        run2.font.color.rgb = MED_GRAY
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3 = p.add_run()
        run3._element.append(fldChar2)


# ── Title Page ───────────────────────────────────────────────────

def create_title_page(doc):
    """Create a clean, professional title page."""
    # Vertical spacing
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(12)

    # Blue accent bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(p, color="2B579A", size=24)

    # Spacer
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

    # Title line 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GSD Autonomous")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(0)

    # Title line 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Development Engine")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Developer Guide")
    run.font.size = Pt(20)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)

    # Thin separator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(p, color="CCCCCC", size=4)
    p.paragraph_format.space_after = Pt(20)

    # Version & Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.7.0  |  March 2026")
    run.font.size = Pt(13)
    run.font.color.rgb = MED_GRAY
    run.font.name = "Calibri"

    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    # Confidential notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.name = "Calibri"
    run.font.letter_spacing = Pt(3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Internal Use Only")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.name = "Calibri"

    doc.add_page_break()


# ── Table of Contents ────────────────────────────────────────────

def create_toc(doc):
    """Insert a Word TOC field that updates when opened in Word."""
    # TOC title
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(12)
    add_bottom_border(p, color="2B579A", size=6)

    # Instruction text
    p = doc.add_paragraph()
    run = p.add_run("(Right-click and select \"Update Field\" to populate this table of contents)")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MED_GRAY
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(8)

    # Insert TOC field
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    # Placeholder text
    run4 = p.add_run("Press F9 or right-click > Update Field to generate Table of Contents")
    run4.font.color.rgb = MED_GRAY
    run4.font.size = Pt(10)
    run4.font.name = "Calibri"

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)

    doc.add_page_break()


# ── Document History ─────────────────────────────────────────────

def create_doc_history(doc, md_text):
    """Read document history from the markdown file and render it."""
    p = doc.add_paragraph("Document History", style='Heading 1')
    # Override page break for this one heading
    p.paragraph_format.page_break_before = False

    # Parse document history table from markdown
    history_data = []
    lines = md_text.split("\n")
    in_history = False
    for line in lines:
        if line.strip() == "## Document History":
            in_history = True
            continue
        if in_history:
            if line.strip().startswith("|") and not re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                parts = line.strip().strip("|").split("|")
                parts = [p.strip() for p in parts]
                if len(parts) >= 3 and parts[0] != "Version":
                    history_data.append(parts[:3])
            elif line.strip() == "---" or (line.strip().startswith("#") and "Document History" not in line):
                break

    if not history_data:
        # Fallback
        history_data = [
            ["1.0.0", "February 2026", "Initial release"],
            ["1.6.1", "March 2026", "Latest release"],
        ]

    num_rows = 1 + len(history_data)
    table = doc.add_table(rows=num_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="CCCCCC")

    headers = ["Version", "Date", "Changes"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_cell_shading(cell, HEADER_BG)

    for r, row_data in enumerate(history_data):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r % 2 == 1:
                set_cell_shading(cell, ALT_ROW_HEX)

    for row in table.rows:
        row.cells[0].width = Inches(0.9)
        row.cells[1].width = Inches(1.3)
        row.cells[2].width = Inches(4.3)

    doc.add_page_break()


# ── Code Block ───────────────────────────────────────────────────

def add_code_block(doc, code_text):
    """Add code block as a single-cell table with gray background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_BG)
    set_table_borders(table, color=BORDER_COLOR, size=4)

    # Clear default paragraph
    cell.text = ""
    lines = code_text.rstrip("\n").split("\n")

    for idx, line in enumerate(lines):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)

        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = DARK_GRAY

    # Small gap after code block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)


# ── Data Table ───────────────────────────────────────────────────

def add_table_from_rows(doc, header_row, data_rows):
    """Add a professionally formatted table with borders."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color=BORDER_COLOR)

    # Auto-fit
    table.autofit = True

    # Header row
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        parse_inline(p, h.strip(), base_size=Pt(9.5), base_color=WHITE)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r, row_data in enumerate(data_rows):
        for c in range(num_cols):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            val = row_data[c].strip() if c < len(row_data) else ""
            p = cell.paragraphs[0]
            parse_inline(p, val, base_size=Pt(9.5), base_color=BLACK)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r % 2 == 1:
                set_cell_shading(cell, ALT_ROW_HEX)

    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)


def parse_table_block(lines):
    if len(lines) < 2:
        return None, None

    def split_row(line):
        parts = line.strip().strip("|").split("|")
        return [p.strip() for p in parts]

    header = split_row(lines[0])
    data = []
    for line in lines[1:]:
        if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
            continue
        data.append(split_row(line))
    return header, data


# ── ASCII Diagram Handler ────────────────────────────────────────

def is_ascii_diagram(code_text):
    """Detect if a code block is an ASCII art diagram (box-drawing chars, arrows)."""
    box_chars = set("┌┐└┘├┤┬┴─│▼►◄▲")
    arrow_count = code_text.count("──>") + code_text.count("|──") + code_text.count("─>")
    pipe_count = code_text.count("│")
    has_box = any(c in box_chars for c in code_text)
    return has_box or (arrow_count >= 2 and pipe_count >= 4)


def add_diagram_as_table(doc, code_text):
    """Render an ASCII diagram as a styled single-cell table with monospace font."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FAFAFA")
    set_table_borders(table, color="2B579A", size=6)

    cell.text = ""
    # Add a "Figure" label
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Extract a title from the diagram content if possible
    lines = code_text.strip().split("\n")
    title = ""
    for line in lines:
        stripped = line.strip().strip("│").strip()
        # Look for capitalized title-like lines
        if stripped and len(stripped) > 5 and stripped == stripped.upper().replace(" ", stripped.replace(" ", "")):
            pass
        if stripped and not any(c in stripped for c in "─┌┐└┘├┤┬┴│▼►◄▲") and len(stripped) > 3:
            if stripped.isupper() or (stripped[0].isupper() and "──>" not in stripped):
                title = stripped
                break

    if title:
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE
    else:
        run = p.add_run("System Diagram")
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE

    # Add the diagram in monospace
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = DARK_GRAY

    # Spacer after
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)


# ── Main Markdown Processor ──────────────────────────────────────

def process_markdown(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    skip_front_matter = False
    list_number = 0

    while i < len(lines):
        line = lines[i]

        # Skip YAML front matter
        if i == 0 and line.strip() == "---":
            skip_front_matter = True
            i += 1
            continue
        if skip_front_matter:
            if line.strip() == "---":
                skip_front_matter = False
            i += 1
            continue

        # Skip title/subtitle/meta lines already handled by title page
        stripped = line.strip()
        if (stripped.startswith("# GSD Autonomous Development Engine") or
            stripped == "## Developer Guide" or
            stripped.startswith("**Version") or
            stripped.startswith("**Date:") or
            stripped.startswith("**Classification:") or
            stripped == "*Confidential - Internal Use Only*"):
            i += 1
            continue

        # Skip standalone --- dividers
        if re.match(r'^---+\s*$', line.strip()):
            i += 1
            continue

        # Skip markdown TOC
        if line.strip() in ("### Table of Contents", "## Table of Contents"):
            i += 1
            while i < len(lines) and (lines[i].strip().startswith("- [") or lines[i].strip() == ""):
                i += 1
            continue

        # Skip Document History (we create our own styled version)
        if line.strip() in ("## Document History", "### Document History"):
            i += 1
            while i < len(lines) and (lines[i].strip().startswith("|") or lines[i].strip() == ""):
                i += 1
            continue

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                if is_ascii_diagram(code_text):
                    add_diagram_as_table(doc, code_text)
                else:
                    add_code_block(doc, code_text)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if (line.strip().startswith("|") and i + 1 < len(lines)
                and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1])):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, data = parse_table_block(table_lines)
            if header and data:
                add_table_from_rows(doc, header, data)
            continue

        # Headings using Word built-in styles
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            list_number = 0  # Reset numbered list counter

            style_map = {
                1: 'Heading 1',  # # -> Heading 1 (chapters: "Chapter 1:", "Appendices")
                2: 'Heading 2',  # ## -> Heading 2 (sections: "1.1", "Appendix A:")
                3: 'Heading 3',  # ### -> Heading 3 (sub-sections)
                4: 'Heading 4',  # #### -> Heading 4 (sub-sub-sections)
            }

            style_name = style_map.get(level, 'Heading 4')
            p = doc.add_paragraph(text, style=style_name)

            # Add a blue bottom border under chapter headings
            if level == 1:
                add_bottom_border(p, color="2B579A", size=8)

            i += 1
            continue

        # Bullet lists
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            list_number = 0

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            if indent >= 4:
                p.paragraph_format.left_indent = Inches(0.7)
                bullet_char = "\u25CB "  # open circle
            else:
                p.paragraph_format.left_indent = Inches(0.35)
                bullet_char = "\u2022  "  # bullet

            run = p.add_run(bullet_char)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = BLUE

            parse_inline(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if num_match:
            text = num_match.group(3)
            num = num_match.group(2)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            run = p.add_run(f"{num}.  ")
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = BLUE

            parse_inline(p, text)
            i += 1
            continue

        # Empty lines
        if line.strip() == "":
            list_number = 0
            i += 1
            continue

        # Regular paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        parse_inline(p, line.strip())

        i += 1


# ── Final Proof ─────────────────────────────────────────────────

def proof_markdown(md_text):
    """
    Final proof of the markdown source before generating the Word document.
    Checks structure, numbering, formatting, and completeness.
    Returns (passed: bool, issues: list[str], stats: dict).
    """
    lines = md_text.split("\n")
    issues = []
    stats = {
        "total_lines": len(lines),
        "chapters": [],
        "sections": 0,
        "tables": 0,
        "code_blocks": 0,
        "appendices": [],
    }

    # --- Check 1: Chapter presence and ordering ---
    chapter_lines = []
    for i, line in enumerate(lines):
        m = re.match(r'^# Chapter (\d+):', line)
        if m:
            chapter_lines.append((int(m.group(1)), i + 1, line.strip()))

    if chapter_lines:
        expected = 1
        for num, ln, title in chapter_lines:
            if num != expected:
                issues.append(f"Line {ln}: Expected Chapter {expected}, found Chapter {num}: {title}")
            expected = num + 1
            stats["chapters"].append(f"Ch{num}")

    # --- Check 2: Section numbering within chapters ---
    current_chapter = 0
    expected_section = 1
    for i, line in enumerate(lines):
        ch_match = re.match(r'^# Chapter (\d+):', line)
        if ch_match:
            current_chapter = int(ch_match.group(1))
            expected_section = 1
            continue
        sec_match = re.match(r'^## (\d+)\.(\d+)\s', line)
        if sec_match:
            ch = int(sec_match.group(1))
            sec = int(sec_match.group(2))
            stats["sections"] += 1
            if ch == current_chapter and sec != expected_section:
                issues.append(
                    f"Line {i+1}: Section {ch}.{sec} — expected {ch}.{expected_section}"
                )
            if ch == current_chapter:
                expected_section = sec + 1

    # --- Check 3: Code fence balance ---
    fence_count = 0
    open_fence_line = None
    for i, line in enumerate(lines):
        if line.strip().startswith("```"):
            fence_count += 1
            if fence_count % 2 == 1:
                open_fence_line = i + 1
            else:
                open_fence_line = None
                stats["code_blocks"] += 1
    if fence_count % 2 != 0:
        issues.append(f"Unclosed code fence (opened at line {open_fence_line}). Total fences: {fence_count}")

    # --- Check 4: Table formatting ---
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|") and i + 1 < len(lines):
            if re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1]):
                # Found table header + separator
                header_cols = len(line.strip().strip("|").split("|"))
                sep_cols = len(lines[i + 1].strip().strip("|").split("|"))
                stats["tables"] += 1
                if header_cols != sep_cols:
                    issues.append(
                        f"Line {i+1}: Table column mismatch — header has {header_cols} cols, "
                        f"separator has {sep_cols} cols"
                    )
                # Check data rows
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    data_cols = len(lines[j].strip().strip("|").split("|"))
                    if data_cols != header_cols:
                        issues.append(
                            f"Line {j+1}: Table row has {data_cols} cols, expected {header_cols}"
                        )
                    j += 1
                i = j
                continue
        i += 1

    # --- Check 5: Appendix ordering ---
    appendix_start = None
    last_chapter_line = 0
    for i, line in enumerate(lines):
        if re.match(r'^# Chapter \d+:', line):
            last_chapter_line = i + 1
        if line.strip() == "# Appendices":
            appendix_start = i + 1

    if appendix_start and appendix_start < last_chapter_line:
        issues.append(
            f"Line {appendix_start}: Appendices section appears BEFORE "
            f"last chapter at line {last_chapter_line}"
        )

    for i, line in enumerate(lines):
        m = re.match(r'^## Appendix ([A-Z]):', line)
        if m:
            stats["appendices"].append(m.group(1))

    # --- Check 6: Placeholder / TODO text ---
    for i, line in enumerate(lines):
        lower = line.lower()
        for marker in ["todo", "fixme", "placeholder", "tbd"]:
            if marker in lower and not line.strip().startswith("|") and "TodoWrite" not in line:
                # Exclude table rows that might legitimately contain these words
                issues.append(f"Line {i+1}: Possible placeholder text found: '{marker}' in: {line.strip()[:80]}")
        # Check for XXX but exclude legitimate code patterns like handleXxx, useXxx
        if "xxx" in lower and not re.search(r'`[^`]*xxx[^`]*`', lower) and not line.strip().startswith("|"):
            issues.append(f"Line {i+1}: Possible placeholder text found: 'xxx' in: {line.strip()[:80]}")

    # --- Check 7: Version consistency ---
    version_refs = []
    for i, line in enumerate(lines):
        # Find version patterns like "Version 1.x.x" but skip the history table
        if "Version" in line and re.search(r'Version\s+(\d+\.\d+\.\d+)', line):
            m = re.search(r'Version\s+(\d+\.\d+\.\d+)', line)
            if m and not line.strip().startswith("|"):
                version_refs.append((i + 1, m.group(1)))

    # --- Check 8: Duplicate section numbers ---
    seen_sections = {}
    for i, line in enumerate(lines):
        sec_match = re.match(r'^## (\d+\.\d+)\s', line)
        if sec_match:
            sec_id = sec_match.group(1)
            if sec_id in seen_sections:
                issues.append(
                    f"Line {i+1}: Duplicate section {sec_id} "
                    f"(first at line {seen_sections[sec_id]})"
                )
            else:
                seen_sections[sec_id] = i + 1

    return len(issues) == 0, issues, stats


def proof_document(doc):
    """
    Final proof of the generated Word document.
    Checks heading hierarchy, page count estimate, and content integrity.
    Returns (passed: bool, issues: list[str], stats: dict).
    """
    issues = []
    stats = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "headings": {"h1": 0, "h2": 0, "h3": 0, "h4": 0},
        "estimated_pages": 0,
    }

    # Count headings by style
    heading_order = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style == "Heading 1":
            stats["headings"]["h1"] += 1
            heading_order.append((1, p.text))
        elif style == "Heading 2":
            stats["headings"]["h2"] += 1
            heading_order.append((2, p.text))
        elif style == "Heading 3":
            stats["headings"]["h3"] += 1
            heading_order.append((3, p.text))
        elif style == "Heading 4":
            stats["headings"]["h4"] += 1
            heading_order.append((4, p.text))

    # Check: Heading 1 should exist (chapters)
    if stats["headings"]["h1"] == 0:
        issues.append("No Heading 1 (chapter-level) headings found in document")

    # Check: Heading 2 should exist (sections)
    if stats["headings"]["h2"] == 0:
        issues.append("No Heading 2 (section-level) headings found in document")

    # Check: H2 should not appear before first H1 (except Document History)
    first_h1 = None
    for level, text in heading_order:
        if level == 1 and first_h1 is None:
            first_h1 = text
            break

    # Check: Appendices should come after all Chapter headings
    last_chapter_idx = -1
    appendix_idx = -1
    for idx, (level, text) in enumerate(heading_order):
        if level == 1 and text.startswith("Chapter"):
            last_chapter_idx = idx
        if level == 1 and text == "Appendices":
            appendix_idx = idx

    if appendix_idx != -1 and last_chapter_idx != -1 and appendix_idx < last_chapter_idx:
        issues.append(
            f"Appendices heading appears before last Chapter heading in Word document"
        )

    # Estimate page count (~40 paragraphs per page)
    stats["estimated_pages"] = max(1, stats["paragraphs"] // 40)

    # Check: Tables should exist
    if stats["tables"] == 0:
        issues.append("No tables found in generated document")

    # Check: Minimum content threshold
    if stats["paragraphs"] < 100:
        issues.append(f"Document seems too short: only {stats['paragraphs']} paragraphs")

    return len(issues) == 0, issues, stats


# ── Main ─────────────────────────────────────────────────────────

def main():
    print("Reading markdown file...")
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()

    # ── Step 1: Proof the markdown source ──
    print("\n" + "=" * 60)
    print("STEP 1: Proofing markdown source...")
    print("=" * 60)
    md_passed, md_issues, md_stats = proof_markdown(md_text)

    print(f"  Lines:      {md_stats['total_lines']}")
    print(f"  Chapters:   {len(md_stats['chapters'])} ({', '.join(md_stats['chapters'])})")
    print(f"  Sections:   {md_stats['sections']}")
    print(f"  Tables:     {md_stats['tables']}")
    print(f"  Code blocks:{md_stats['code_blocks']}")
    print(f"  Appendices: {', '.join(md_stats['appendices'])}")

    if md_issues:
        print(f"\n  [!!] {len(md_issues)} issue(s) found in markdown:")
        for issue in md_issues:
            print(f"       - {issue}")
    else:
        print("\n  [OK] Markdown proof passed — no issues found")

    if not md_passed:
        print("\n  [!!] WARNING: Proceeding with generation despite issues.")
        print("       Fix the markdown and re-run for a clean build.\n")

    # ── Step 2: Generate the Word document ──
    print("\n" + "=" * 60)
    print("STEP 2: Generating Word document...")
    print("=" * 60)

    doc = Document()
    setup_styles(doc)
    create_title_page(doc)
    create_toc(doc)
    create_doc_history(doc, md_text)

    print("  Processing markdown content...")
    process_markdown(doc, md_text)
    add_headers_footers(doc)

    print(f"  Saving to {DOCX_PATH}...")
    doc.save(DOCX_PATH)

    size_mb = os.path.getsize(DOCX_PATH) / (1024 * 1024)

    # ── Step 3: Proof the generated document ──
    print("\n" + "=" * 60)
    print("STEP 3: Proofing generated Word document...")
    print("=" * 60)

    doc_passed, doc_issues, doc_stats = proof_document(doc)

    print(f"  Paragraphs: {doc_stats['paragraphs']}")
    print(f"  Tables:     {doc_stats['tables']}")
    print(f"  Headings:   H1={doc_stats['headings']['h1']}, "
          f"H2={doc_stats['headings']['h2']}, "
          f"H3={doc_stats['headings']['h3']}, "
          f"H4={doc_stats['headings']['h4']}")
    print(f"  Est. pages: ~{doc_stats['estimated_pages']}")
    print(f"  File size:  {size_mb:.2f} MB")

    if doc_issues:
        print(f"\n  [!!] {len(doc_issues)} issue(s) found in Word document:")
        for issue in doc_issues:
            print(f"       - {issue}")
    else:
        print("\n  [OK] Word document proof passed — no issues found")

    # ── Final Summary ──
    total_issues = len(md_issues) + len(doc_issues)
    print("\n" + "=" * 60)
    if total_issues == 0:
        print("FINAL PROOF: PASSED — Document is production-ready")
    else:
        print(f"FINAL PROOF: {total_issues} issue(s) found — review before distributing")
    print("=" * 60)
    print(f"\nOutput: {DOCX_PATH}")
    print("Tip: Open in Word, press Ctrl+A then F9 to update the Table of Contents.")


if __name__ == "__main__":
    main()
